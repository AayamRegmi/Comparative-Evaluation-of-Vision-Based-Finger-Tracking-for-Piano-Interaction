"""
Builds Documentation/Testing and Evaluation.docx
Pulls real screenshots from the main report (extracted to data/extracted_imgs/).
Run with:  piano_env\Scripts\python.exe build_testing_evaluation.py
"""

from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image, ImageDraw, ImageFont

OUT_PATH  = Path("Documentation/Testing and Evaluation.docx")
IMG_DIR   = Path("data/extracted_imgs")

COL_HEADING = RGBColor(0x1F, 0x39, 0x64)

# ---------------------------------------------------------------------------
# Image name map: descriptive key -> filename in data/extracted_imgs/
# ---------------------------------------------------------------------------
IMGS = {
    "fitzpatrick_scale":      "p287_rId22.png",
    "lux_meter":              "p293_rId24.png",
    "mediapipe_preview":      "p309_rId25.png",
    "polygon_mask":           "p315_rId26.png",
    "polygon_preview":        "p317_rId28.png",
    "polygon_controls":       "p319_rId29.png",
    "homography_mask":        "p332_rId32.png",
    "correct_detection":      "p335_rId33.png",
    "false_detection":        "p336_rId34.png",
    "vertical_depth_viz":     "p338_rId35.png",
    "live_stats":             "p347_rId36.png",
    "live_results_L":         "p355_rId41.png",
    "live_results_R":         "p355_rId42.png",
    "raw_data_structure":     "p368_rId43.png",
    "benchmark_app":          "p382_rId44.png",
    "data_diversity":         "p388_rId45.png",
    "detection_rate_chart":   "p397_rId46.png",
    "per_hand_mediapipe":     "p404_rId47.png",
    "per_hand_openpose":      "p408_rId48.png",
    "dfr_skin_tone":          "p413_rId49.png",
    "dfr_lighting":           "p422_rId50.png",
    "mpjpe_comparison":       "p437_rId51.png",
    "mask_misalignment":      "p444_rId52.png",
    "per_finger_mpjpe_combined": "p448_rId53.png",
    "per_finger_mpjpe_mp":    "p450_rId54.png",
    "mp_hitrate":             "p452_rId47.png",
    "openpose_mpjpe":         "p457_rId55.png",
    "fitzpatrick_mpjpe":      "p476_rId56.png",
    "mpjpe_by_lux":           "p489_rId57.png",
    "mpjpe_vs_lux_scatter":   "p491_rId58.png",
    "mpjpe_vs_handsize":      "p503_rId59.png",
}

# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shd(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _cell_shd(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _left_bar(paragraph, fill_hex, bar_color="1F3964"):
    _shd(paragraph, fill_hex)
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), bar_color)
    pBdr.append(left)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = COL_HEADING
    return p


def add_body(doc, text):
    return doc.add_paragraph(text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_note(doc, text, kind="note"):
    colours = {
        "note":    ("FFF3CD", "856404", "7B6000"),
        "info":    ("D1ECF1", "0C5460", "0C5460"),
        "success": ("D6EAD6", "1A5C2A", "1A5C2A"),
        "warning": ("F8D7DA", "842029", "CC0000"),
    }
    bg, left, txt_hex = colours.get(kind, colours["note"])
    prefix = {"note": "NOTE  ", "info": "FINDING  ",
              "success": "RESULT  ", "warning": "LIMITATION  "}.get(kind, "NOTE  ")
    p  = doc.add_paragraph()
    rb = p.add_run(prefix)
    rb.bold = True
    rb.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    rt = p.add_run(text)
    rt.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    _left_bar(p, bg, left)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_image(doc, key, caption, width_in=5.5):
    """Insert a real screenshot if available, else a labelled placeholder."""
    fpath = IMG_DIR / IMGS.get(key, "")
    if fpath.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(fpath), width=Inches(width_in))
    else:
        # Fallback grey placeholder
        w_px, h_px = int(width_in * 96), 260
        img  = Image.new("RGB", (w_px, h_px), (180, 185, 195))
        draw = ImageDraw.Draw(img)
        draw.rectangle([3, 3, w_px-4, h_px-4], outline=(130,135,145), width=3)
        try:
            fnt = ImageFont.truetype("arialbd.ttf", 16)
        except OSError:
            fnt = ImageFont.load_default()
        txt = f"[ {caption} ]"
        bb  = draw.textbbox((0,0), txt, font=fnt)
        draw.text(((w_px-(bb[2]-bb[0]))//2, (h_px-16)//2), txt,
                  font=fnt, fill=(80,85,95))
        buf = BytesIO()
        img.save(buf, "PNG")
        buf.seek(0)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(buf, width=Inches(width_in))

    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size   = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap.paragraph_format.space_after = Pt(8)


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        r.font.size = Pt(10)
        _cell_shd(c, "1F3964")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri+1]
        bg  = "FFFFFF" if ri % 2 == 0 else "F5F7FA"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            _cell_shd(c, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Testing and Evaluation", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.color.rgb = COL_HEADING

    sub = doc.add_paragraph(
        "Comparative Evaluation of Vision-Based Finger Tracking for Piano Interaction"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size   = Pt(12)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── 1. Overview ──────────────────────────────────────────────────────────
    add_heading(doc, "1.  Overview", 1)
    add_body(doc,
        "This document covers the testing and validation of the data collection "
        "pipeline and the evaluation of research results.  It is organised into "
        "three parts:")
    add_bullet(doc,
        "Part A  (Sections 2–4):  Feature validation — testing each pipeline "
        "component before data collection began.")
    add_bullet(doc,
        "Part B  (Section 5):  Requirements evaluation — formal completion status "
        "of all functional and non-functional requirements.")
    add_bullet(doc,
        "Part C  (Sections 6–10):  Research evaluation — analysis of results "
        "across 38 participant sessions.")
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PART A — FEATURE VALIDATION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART A — Feature Validation", 1)
    add_body(doc,
        "Before data collection began, each pipeline feature was individually "
        "tested to verify correct behaviour and eliminate systematic bias in the "
        "MPJPE measurements.")
    doc.add_paragraph()

    # ── 2. Environmental characterisation ────────────────────────────────────
    add_heading(doc, "2.  Environmental and Participant Characterisation", 1)

    add_heading(doc, "2.1  Fitzpatrick Skin Type Auto-Detection", 2)
    add_body(doc,
        "Skin tone classification was implemented using the ITA (Individual "
        "Typology Angle) colorimetric method (Chardon et al., 1991).  "
        "The system samples CIELab pixel values from the palm convex hull "
        "(landmarks 0, 1, 5, 9, 13, 17) across 20 frames to estimate the "
        "participant's Fitzpatrick type automatically, removing the need for "
        "manual classification and eliminating observer bias.")
    doc.add_paragraph()
    add_body(doc, "ITA formula:")
    add_body(doc, "     ITA = arctan( (L* - 50) / b* )  x  (180 / pi)")
    add_body(doc,
        "Where L* is perceptual lightness and b* is the yellow-blue opponent "
        "channel.  The resulting angle maps to Fitzpatrick types I-VI using "
        "the Chardon (1991) thresholds.")
    doc.add_paragraph()
    add_image(doc, "fitzpatrick_scale",
              "Fig 6.2.1.a: Fitzpatrick Scale reference chart (Fitzpatrick, 1988)", width_in=4.5)

    add_heading(doc, "2.2  Ambient Illuminance Classification", 2)
    add_body(doc,
        "Ambient lux was entered manually by the researcher at the start of each "
        "session using a mobile phone lux meter app.  Manual entry was preferred "
        "over automatic camera-based measurement to avoid introducing any "
        "measurement-induced variation across sessions.")
    add_body(doc, "Lux thresholds (defined in config.py):")
    add_table(doc,
        ["Category", "Lux range"],
        [["Dim",     "< 100 lux"],
         ["Indoor",  "100 – 499 lux"],
         ["Bright",  ">= 500 lux"]],
        col_widths=[1.5, 2.5],
    )
    add_image(doc, "lux_meter",
              "Fig 6.2.2.a: Mobile phone lux meter showing 14.67 lux reading", width_in=3.5)

    add_heading(doc, "2.3  Hand Size Classification", 2)
    add_body(doc,
        "Hand size was measured manually before each session using a tape measure "
        "from the base of the palm to the tip of the middle finger with the hand "
        "spread flat.  Values ranged from 7.7 cm to 12.2 cm across all participants.  "
        "Size was recorded as a continuous variable rather than binned categories "
        "to preserve regression analysis precision.")
    doc.add_paragraph()

    # ── 3. MPJPE calibration and bias reduction ──────────────────────────────
    add_heading(doc, "3.  MPJPE Calibration and Bias Reduction", 1)

    add_heading(doc, "3.1  Horizontal-Only MPJPE", 2)
    add_body(doc,
        "The initial 2D MPJPE implementation used full Euclidean distance from "
        "fingertip to key centre.  This introduced systematic bias because fingers "
        "at rest sit higher in the frame than fingers actively pressing keys — "
        "a natural anatomical variation that penalised correct detections.")
    add_body(doc,
        "The metric was revised to measure horizontal displacement only:")
    add_body(doc, "     h_err = | tip_x  -  key_cx |")
    add_body(doc,
        "This removes the vertical (depth/anatomy) component entirely.  "
        "Any fingertip Y position within the key height is treated as correct, "
        "and only horizontal misalignment contributes to the error score.")
    doc.add_paragraph()
    add_image(doc, "mediapipe_preview",
              "Fig 6.3.1.a: MediaPipe hand landmark preview showing fingertip positions over keys",
              width_in=5.5)

    add_heading(doc, "3.2  Polygon Containment for Finger Detection", 2)
    add_body(doc,
        "A key polygon is a warped quadrilateral matching the physical boundaries "
        "of each piano key in the camera view.  "
        "For each note_on event, a pointPolygonTest determines whether any detected "
        "fingertip falls inside the triggered key's polygon.  "
        "This replaced an earlier centroid-proximity approach that "
        "misclassified fingers near adjacent keys.")
    doc.add_paragraph()
    add_image(doc, "polygon_mask",
              "Fig 6.3.2.b: Polygon containment mask — key boundaries drawn on camera frame",
              width_in=5.5)
    add_image(doc, "polygon_preview",
              "Fig 6.3.2.c: Polygon containment mask preview — fingertips coloured by key match",
              width_in=5.5)
    add_image(doc, "polygon_controls",
              "Fig 6.3.2.d: Polygon containment mask control panel", width_in=5.0)

    add_heading(doc, "3.3  Perspective Homography Calibration", 2)
    add_body(doc,
        "The overhead tripod position produces an oblique camera angle that distorts "
        "the apparent shape of piano keys.  A homography transform was applied to "
        "warp the key grid from flat 2D space to the camera's projected perspective, "
        "ensuring the mask polygons align with the visible key boundaries even at "
        "steep viewing angles.")
    add_body(doc, "Homography model:")
    add_body(doc, "     [x', y', w']^T  =  H  [x, y, 1]^T")
    add_body(doc,
        "where H is a 3x3 perspective transform matrix computed from the four "
        "corner drag handles in the calibration UI.  "
        "Warped polygon vertices are recovered by dehomogenisation: "
        "x_screen = x'/w',  y_screen = y'/w'.")
    doc.add_paragraph()
    add_image(doc, "homography_mask",
              "Fig 6.3.3.a: Homography calibration mask — warped overlay aligned to oblique camera view",
              width_in=5.5)
    add_image(doc, "correct_detection",
              "Fig 6.3.3.b: Correct fingertip detection — tip inside warped key polygon",
              width_in=5.5)
    add_image(doc, "false_detection",
              "Fig 6.3.3.c: Out-of-polygon false detection — tip outside key boundary",
              width_in=5.5)
    add_image(doc, "vertical_depth_viz",
              "Fig 6.3.3.d: Full-depth vertical line visualisation — shows that any Y within key height is valid",
              width_in=5.5)

    # ── 4. Pipeline feature tests ─────────────────────────────────────────────
    add_heading(doc, "4.  Pipeline Feature Tests", 1)

    add_heading(doc, "4.1  MIDI–Video Temporal Alignment", 2)
    add_body(doc,
        "MIDI and video timestamps were aligned using a shared time origin "
        "(rec_t0) set at the moment the SPACE key is pressed to begin recording.  "
        "The MIDI recorder runs as a background daemon thread, polling the MIDI "
        "port every 1 ms using mido's iter_pending() interface.  "
        "Each MIDI event is timestamped as (time.perf_counter() - rec_t0).  "
        "The FrameLogger records (frame_index, time.perf_counter() - rec_t0) "
        "once per writer.write(frame) call.  "
        "Analyse.py performs a binary-search lookup to find the video frame "
        "nearest in time to each MIDI note_on event.")
    doc.add_paragraph()
    add_note(doc,
        "Sub-frame temporal precision:  at 30 FPS, one frame = 33 ms.  "
        "The MIDI poll interval of 1 ms gives < 3% temporal quantisation error "
        "relative to the frame window.",
        kind="info")

    add_heading(doc, "4.2  System Performance Statistics", 2)
    add_body(doc,
        "FPS, MediaPipe inference latency, and total frame latency were monitored "
        "in real time during each session using the StatsCollector class.  "
        "A warmup period (configurable in config.py) was excluded from rolling "
        "averages to allow the camera and model to stabilise.")
    doc.add_paragraph()
    add_image(doc, "live_stats",
              "Fig 6.5.a: live_preview.py performance stats overlay — FPS, inference, and latency",
              width_in=5.5)

    add_heading(doc, "4.3  Live MPJPE Test Results", 2)
    add_body(doc,
        "live_preview.py displays a real-time results panel after each test session.  "
        "The panel shows per-hand and per-finger MPJPE, accuracy percentage, "
        "and detection rate.")
    doc.add_paragraph()
    add_image(doc, "live_results_L",
              "Fig 6.5.b: live_preview.py end-of-session results panel — left hand breakdown",
              width_in=5.5)
    add_image(doc, "live_results_R",
              "Fig 6.5.b: live_preview.py end-of-session results panel — right hand breakdown",
              width_in=5.5)

    add_heading(doc, "4.4  Raw Data Storage Structure", 2)
    add_body(doc,
        "Each session writes to a dedicated data/raw/p###/ folder containing the "
        "raw video, frame timestamp CSV, MIDI JSONL, standard MIDI file, and a "
        "session metadata JSON with all demographic and environmental variables.")
    doc.add_paragraph()
    add_image(doc, "raw_data_structure",
              "Fig 6.3.3.a: Raw data folder structure for one participant session",
              width_in=5.0)

    # ════════════════════════════════════════════════════════════════════════
    # PART B — REQUIREMENTS EVALUATION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART B — Requirements Evaluation", 1)
    doc.add_paragraph()

    # ── 5. Requirements completion ────────────────────────────────────────────
    add_heading(doc, "5.  Requirements Completion", 1)

    add_heading(doc, "5.1  Functional Requirements", 2)
    add_table(doc,
        ["#", "Requirement", "MoSCoW", "Status"],
        [
            ["FR-01", "Finger landmark tracking via vision models",   "Must",   "Completed"],
            ["FR-02", "MIDI ground truth capture",                    "Must",   "Completed"],
            ["FR-03", "MIDI-video temporal synchronisation",          "Must",   "Completed"],
            ["FR-04", "Multiple open-source model support",           "Must",   "Partially completed  (MoveNet excluded)"],
            ["FR-05", "Camera-to-keyboard spatial calibration",       "Must",   "Completed  (homography polygon mask)"],
            ["FR-06", "Data logging system",                          "Must",   "Completed  (CSV, JSONL, JSON, MP4)"],
            ["FR-07", "Keypress event analysis",                      "Must",   "Completed  (polygon containment)"],
            ["FR-08", "Quantitative performance metrics",             "Must",   "Completed  (MPJPE, detection rate, FPS)"],
            ["FR-09", "Controlled variation of conditions",           "Must",   "Completed  (Fitzpatrick, lux, hand size)"],
            ["FR-10", "Statistical model comparison",                 "Must",   "Completed  (descriptive stats + charts)"],
            ["FR-11", "Model selection justification",                "Must",   "Completed  (report sections 3 and 8)"],
            ["FR-12", "Reproducible protocol documentation",          "Should", "Completed"],
            ["FR-13", "Additional model evaluation",                  "Could",  "Not completed  (MoveNet format mismatch)"],
            ["FR-14", "Real-time visualisation overlay",              "Should", "Completed  (live_preview.py)"],
            ["FR-15", "Automatic occlusion annotation",               "Could",  "Not completed  (descoped)"],
            ["FR-16", "3D hand pose estimation",                      "Won't",  "Correctly excluded"],
            ["FR-17", "Mobile / embedded deployment",                 "Won't",  "Correctly excluded"],
            ["FR-18", "Multi-camera tracking",                        "Won't",  "Correctly excluded"],
            ["FR-19", "Real-time audio feedback",                     "Won't",  "Correctly excluded"],
        ],
        col_widths=[0.65, 2.4, 0.75, 3.5],
    )

    add_heading(doc, "5.2  Non-Functional Requirements", 2)
    add_table(doc,
        ["#", "Requirement", "MoSCoW", "Status"],
        [
            ["NFR-01", "Research-focused prototype",                  "Must",   "Completed"],
            ["NFR-02", "Fair and controlled comparison conditions",   "Must",   "Completed"],
            ["NFR-03", "Python + open-source libraries only",         "Must",   "Completed"],
            ["NFR-04", "Ethical compliance and participant protection","Must",   "Completed  (ethics stages 1 and 2)"],
            ["NFR-05", "Real-time performance >= 15-20 FPS",          "Should", "Completed  (30 FPS sustained)"],
            ["NFR-06", "Code quality and reproducibility",            "Should", "Completed"],
            ["NFR-07", "Basic robustness and error handling",         "Could",  "Completed"],
        ],
        col_widths=[0.65, 2.8, 0.75, 3.1],
    )

    add_note(doc,
        "All 11 Must requirements are completed.  "
        "FR-04 is partially completed — MediaPipe and OpenPose are both integrated; "
        "MoveNet was excluded because its output format is incompatible with the "
        "21-keypoint landmark schema used throughout the pipeline.",
        kind="success")

    add_heading(doc, "5.3  Prototype Benchmark Application", 2)
    add_body(doc,
        "A prototype studio-style GUI was developed to allow researchers to switch "
        "between pipeline scripts using on-screen buttons rather than terminal "
        "commands.  This was evaluated as a potential future tool to improve "
        "usability during data collection sessions.")
    doc.add_paragraph()
    add_image(doc, "benchmark_app",
              "Fig 7.3.a: Prototype benchmark application UI design", width_in=5.5)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    # PART C — RESEARCH EVALUATION
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "PART C — Research Evaluation", 1)
    add_body(doc,
        "38 participant sessions were completed.  89,560 video frames, 9,557 MIDI "
        "events, and 4,777 note-on events were processed.  Participants covered "
        "Fitzpatrick types II–V, with the majority in types III and IV.  "
        "All sessions were conducted indoors under Dim or Indoor lighting.")
    doc.add_paragraph()
    add_image(doc, "data_diversity",
              "Fig 8.a: Dataset diversity charts — Fitzpatrick type, lux, and hand size distributions",
              width_in=5.5)

    # ── 6. Detection rate ─────────────────────────────────────────────────────
    add_heading(doc, "6.  Model Detection Rate", 1)

    add_heading(doc, "6.1  Overall Detection Rate", 2)
    add_body(doc,
        "Detection rate is the proportion of note_on events where at least one "
        "fingertip landmark was found inside the triggered key's polygon.")
    add_table(doc,
        ["Outcome", "MediaPipe (n)", "MediaPipe (%)", "OpenPose (n)", "OpenPose (%)"],
        [
            ["Matched",          "3,017", "63.3",  "388",   "8.1"],
            ["Detection fail",   "1,370", "28.7", "1,039",  "21.8"],
            ["Missed",             "390",  "8.2", "3,350",  "70.3"],
            ["Total note-on",    "4,777",  "100", "4,777",  "100"],
        ],
        col_widths=[1.6, 1.2, 1.3, 1.2, 1.3],
    )
    add_image(doc, "detection_rate_chart",
              "Fig 8.1.1.a: MediaPipe vs OpenPose detection rate comparison chart", width_in=5.5)
    add_note(doc,
        "MediaPipe detected landmarks in 91.8% of events (matched + detection_fail).  "
        "OpenPose only detected landmarks in 29.7% of events, because it requires "
        "a tightly cropped hand image as input — the full-scene video feed used "
        "here is not its intended use case.",
        kind="info")

    add_heading(doc, "6.2  Per-Hand and Per-Finger Detection Rate", 2)
    add_body(doc,
        "MediaPipe showed consistent per-finger detection across both hands with "
        "the thumb achieving the highest match rate.  OpenPose detection was too "
        "sparse to produce reliable per-finger statistics.")
    doc.add_paragraph()
    add_image(doc, "per_hand_mediapipe",
              "Fig 8.1.3.a: Per-hand and per-finger detection rate for MediaPipe", width_in=5.5)
    add_image(doc, "per_hand_openpose",
              "Fig 8.1.3.c: Per-finger detection rate for OpenPose (low n — limited reliability)",
              width_in=5.5)

    add_heading(doc, "6.3  Detection-Fail Rate by Skin Tone", 2)
    add_body(doc,
        "Detection-Fail Rate (DFR) measures cases where landmarks were detected "
        "but no tip landed inside the key polygon — indicating positioning error "
        "rather than complete model failure.")
    add_table(doc,
        ["Fitzpatrick Type", "MediaPipe DFR %", "OpenPose DFR %"],
        [["Type II", "16.8", "44.1"],
         ["Type III","32.8", "15.5"],
         ["Type IV", "26.4", "25.2"],
         ["Type V",  "30.5", "31.2"]],
        col_widths=[1.8, 2.0, 2.0],
    )
    add_image(doc, "dfr_skin_tone",
              "Fig 8.1.4.a: Detection-fail rate by Fitzpatrick skin tone type", width_in=5.5)
    add_note(doc,
        "No consistent upward or downward trend with skin tone is visible for "
        "MediaPipe.  The lowest DFR (16.8%) occurs at Type II, and the highest "
        "(32.8%) at Type III — suggesting no systematic skin-tone bias.",
        kind="info")

    add_heading(doc, "6.4  Detection-Fail Rate by Lighting Condition", 2)
    add_table(doc,
        ["Lighting", "MediaPipe DFR %", "OpenPose DFR %"],
        [["Dim (<100 lux)",  "28.6", "21.7"],
         ["Indoor (100-499 lux)", "28.7", "29.9"],
         ["Bright (>=500 lux)", "—", "—"]],
        col_widths=[2.2, 2.0, 2.0],
    )
    add_image(doc, "dfr_lighting",
              "Fig 8.1.5.a: Detection-fail rate by lighting condition", width_in=5.5)
    add_note(doc,
        "MediaPipe showed effectively zero difference between Dim and Indoor "
        "conditions (28.6% vs 28.7%).  Bright conditions were not represented "
        "in the dataset — all sessions were conducted indoors.",
        kind="info")

    # ── 7. MPJPE results ──────────────────────────────────────────────────────
    add_heading(doc, "7.  MPJPE Results", 1)

    add_heading(doc, "7.1  Overall Model Comparison", 2)
    add_body(doc,
        "MPJPE is only computed over matched events (fingertip inside polygon).  "
        "OpenPose's lower MPJPE (5.87 px vs 6.07 px) is calculated on only 8.1% "
        "of total events, making direct comparison misleading — "
        "MediaPipe's result covers 63.3% of all events.")
    add_table(doc,
        ["Metric", "MediaPipe", "OpenPose"],
        [["Mean MPJPE (px)",       "6.07",  "5.87"],
         ["Median MPJPE (px)",     "5.1",   "4.9"],
         ["Detection rate (%)",    "63.3",  "8.1"],
         ["Events matched (n)",    "3,017", "388"]],
        col_widths=[2.5, 2.0, 2.0],
    )
    add_image(doc, "mpjpe_comparison",
              "Fig 8.2.1.a: Overall MPJPE comparison — MediaPipe vs OpenPose", width_in=5.5)

    add_heading(doc, "7.2  2D Mask Perspective Limitation", 2)
    add_body(doc,
        "When the camera is angled overhead, black keys further from the camera "
        "appear narrower and offset relative to their physical position.  "
        "The 2D polygon mask cannot fully compensate for this 3D key geometry.  "
        "Both models were tested against the same mask, so the bias is consistent "
        "and the comparative result remains valid.")
    doc.add_paragraph()
    add_image(doc, "mask_misalignment",
              "Fig 8.2.2.a: Black key 2D calibration mask misalignment due to camera perspective",
              width_in=5.5)

    add_heading(doc, "7.3  MediaPipe Per-Finger MPJPE", 2)
    add_table(doc,
        ["Hand", "Thumb", "Index", "Middle", "Ring", "Pinky"],
        [["Left  (MPJPE px)",  "5.7", "5.8", "6.5", "6.6", "5.5"],
         ["Right (MPJPE px)",  "5.7", "5.8", "6.4", "6.6", "5.3"]],
        col_widths=[1.5, 1.0, 1.0, 1.0, 1.0, 1.0],
    )
    add_image(doc, "per_finger_mpjpe_mp",
              "Fig 8.2.3.b: Per-finger MPJPE breakdown for MediaPipe — Left and Right hands",
              width_in=5.5)
    add_image(doc, "mp_hitrate",
              "Fig 8.2.3.d: MediaPipe per-hand polygon hit rate", width_in=5.5)
    add_note(doc,
        "Ring and Middle fingers consistently show the highest MPJPE across both "
        "hands (6.5-6.6 px), consistent with their anatomical tendency to move "
        "in coordination rather than independently.",
        kind="info")

    add_heading(doc, "7.4  OpenPose Per-Finger MPJPE", 2)
    add_table(doc,
        ["Finger", "Thumb", "Index", "Middle", "Ring", "Pinky"],
        [["MPJPE (px)",  "6.9", "5.8", "6.1", "6.1", "4.4"]],
        col_widths=[1.5, 1.0, 1.0, 1.0, 1.0, 1.0],
    )
    add_image(doc, "openpose_mpjpe",
              "Fig 8.2.4.b: OpenPose per-finger MPJPE — both hands combined", width_in=5.5)
    add_note(doc,
        "OpenPose shows the inverse pattern to MediaPipe: highest MPJPE at the "
        "Thumb (6.9 px), lowest at Pinky (4.4 px).  This may reflect the "
        "different cropping strategy used for OpenPose inference.",
        kind="info")

    # ── 8. Skin tone effect ───────────────────────────────────────────────────
    add_heading(doc, "8.  Effect of Skin Tone on MPJPE", 1)
    add_table(doc,
        ["Fitzpatrick Type", "MediaPipe MPJPE (px)", "OpenPose MPJPE (px)"],
        [["Type II", "5.51", "6.41"],
         ["Type III","6.02", "5.70"],
         ["Type IV", "6.39", "5.53"],
         ["Type V",  "5.96", "6.34"]],
        col_widths=[2.0, 2.2, 2.2],
    )
    add_image(doc, "fitzpatrick_mpjpe",
              "Fig 8.3.2.b: Mean MPJPE per Fitzpatrick type — MediaPipe and OpenPose", width_in=5.5)
    add_note(doc,
        "MediaPipe shows a mild upward trend from Type II (5.51 px) to Type IV "
        "(6.39 px), but the difference is small.  Types I and VI were not "
        "represented in the dataset.  OpenPose shows no consistent trend.",
        kind="info")

    # ── 9. Lighting effect ────────────────────────────────────────────────────
    add_heading(doc, "9.  Effect of Lighting on MPJPE", 1)
    add_table(doc,
        ["Condition", "MediaPipe (px)", "OpenPose (px)"],
        [["Dim (<100 lux)",      "6.03 (SD=1.16, n=28)", "6.17"],
         ["Indoor (100-499 lux)","6.08 (SD=1.01, n=10)", "5.04"]],
        col_widths=[2.0, 2.8, 2.0],
    )
    add_image(doc, "mpjpe_by_lux",
              "Fig 8.4.2.b: Mean MPJPE per lighting category — MediaPipe and OpenPose", width_in=5.5)
    add_image(doc, "mpjpe_vs_lux_scatter",
              "Fig 8.4.2.c: MPJPE vs lux level scatter plot", width_in=5.5)
    add_note(doc,
        "MediaPipe showed no meaningful difference between Dim and Indoor "
        "conditions (6.03 vs 6.08 px).  OpenPose showed more variation "
        "(6.17 vs 5.04 px) but with a small indoor sample (n=10).  "
        "Bright conditions (>=500 lux) were not represented.",
        kind="info")

    # ── 10. Hand size effect ──────────────────────────────────────────────────
    add_heading(doc, "10.  Effect of Hand Size on MPJPE", 1)
    add_table(doc,
        ["Metric", "MediaPipe", "OpenPose"],
        [["Regression slope (px/cm)", "0.162",  "0.336"],
         ["R-squared",                "0.030",  "0.059"],
         ["Hand size range (cm)",     "7.7-12.2", "7.7-12.2"]],
        col_widths=[2.5, 2.0, 2.0],
    )
    add_image(doc, "mpjpe_vs_handsize",
              "Fig 8.5.2.a: MPJPE vs hand size scatter plot with regression lines", width_in=5.5)
    add_note(doc,
        "Both regression R-squared values are below 0.06, indicating that hand "
        "size explains less than 6% of MPJPE variance.  MediaPipe performs "
        "consistently across the full adult hand size range tested without "
        "per-user recalibration.",
        kind="info")

    # ── 11. Summary ───────────────────────────────────────────────────────────
    add_heading(doc, "11.  Summary of Evaluation Findings", 1)

    add_table(doc,
        ["Finding", "MediaPipe", "OpenPose"],
        [
            ["Detection rate",          "63.3%",    "8.1%"],
            ["Mean MPJPE (px)",         "6.07",     "5.87 (on 8.1% of events)"],
            ["Skin tone effect",        "Mild upward trend — not significant", "No consistent trend"],
            ["Lighting effect",         "Negligible (Dim vs Indoor)",          "Larger variation but small n"],
            ["Hand size effect",        "R2=0.03 — negligible",               "R2=0.06 — negligible"],
            ["Practical recommendation","Best for full-scene deployment",      "Requires pre-cropped hand input"],
        ],
        col_widths=[2.2, 2.5, 2.6],
    )

    add_note(doc,
        "MediaPipe is the recommended model for full-scene piano tracking.  "
        "It achieves a mean MPJPE of 6.07 px (well below the ~10 px threshold "
        "for a standard white key) with a detection rate of 63.3%.  "
        "OpenPose's marginally lower MPJPE (5.87 px) is not comparable — "
        "it only detected landmarks on 8.1% of events using the full-scene input.",
        kind="success")

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
