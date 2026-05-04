"""
Builds Documentation/Installation Guidelines.docx from scratch.
Run with:  piano_env\Scripts\python.exe build_install_guide.py
"""

from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

from PIL import Image, ImageDraw, ImageFont

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

OUT_PATH = Path("Documentation/Installation Guidelines.docx")

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

COL_HEADING    = RGBColor(0x1F, 0x39, 0x64)   # dark navy
COL_CODE_BG    = "F2F2F2"                       # hex, light grey
COL_NOTE_BG    = "FFF3CD"                       # amber tint
COL_WARN_BG    = "F8D7DA"                       # red tint
COL_TABLE_HDR  = RGBColor(0x1F, 0x39, 0x64)
COL_PLACEHOLDER = (180, 185, 195)               # PIL RGB grey
COL_PH_BORDER   = (130, 135, 145)
COL_PH_TEXT     = (80, 85, 95)

# ---------------------------------------------------------------------------
# Placeholder image generator
# ---------------------------------------------------------------------------

def make_placeholder(label: str, width_px: int = 600, height_px: int = 340) -> BytesIO:
    """Return a BytesIO PNG of a labelled grey placeholder box."""
    img  = Image.new("RGB", (width_px, height_px), COL_PLACEHOLDER)
    draw = ImageDraw.Draw(img)

    # Border
    bw = 3
    draw.rectangle([bw, bw, width_px - bw - 1, height_px - bw - 1],
                   outline=COL_PH_BORDER, width=bw)

    # Icon lines (simple image icon)
    cx, cy = width_px // 2, height_px // 2 - 28
    icon_w, icon_h = 64, 48
    draw.rectangle([cx - icon_w, cy - icon_h, cx + icon_w, cy + icon_h],
                   outline=COL_PH_TEXT, width=2)
    # Mountain + sun inside icon
    draw.polygon([(cx - icon_w + 8, cy + icon_h - 2),
                  (cx - 10, cy - icon_h + 18),
                  (cx + 28, cy + icon_h - 2)], outline=COL_PH_TEXT)
    draw.polygon([(cx + 10, cy + icon_h - 2),
                  (cx + icon_w - 8, cy - icon_h + 28),
                  (cx + icon_w - 4, cy + icon_h - 2)], outline=COL_PH_TEXT)
    draw.ellipse([cx + 30, cy - icon_h + 4, cx + 52, cy - icon_h + 24],
                 outline=COL_PH_TEXT)

    # Main label
    try:
        font_main  = ImageFont.truetype("arialbd.ttf", 18)
        font_small = ImageFont.truetype("arial.ttf",   14)
    except OSError:
        font_main = font_small = ImageFont.load_default()

    main_text = "[ IMAGE PLACEHOLDER ]"
    bb = draw.textbbox((0, 0), main_text, font=font_main)
    tw = bb[2] - bb[0]
    draw.text(((width_px - tw) // 2, cy + icon_h + 12), main_text,
              font=font_main, fill=COL_PH_TEXT)

    bb2 = draw.textbbox((0, 0), label, font=font_small)
    tw2 = bb2[2] - bb2[0]
    draw.text(((width_px - tw2) // 2, cy + icon_h + 40), label,
              font=font_small, fill=COL_PH_TEXT)

    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _set_shading(paragraph, fill_hex: str):
    """Apply background fill to a paragraph via pPr/shd."""
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _set_cell_shading(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _set_para_border(paragraph, fill_hex: str, left_color: str = "1F3964"):
    """Left-bar + background — used for notes/warnings."""
    _set_shading(paragraph, fill_hex)
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), left_color)
    pBdr.append(left)
    pPr.append(pBdr)


def _cell_no_margins(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    "80")
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = COL_HEADING
    return p


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p


def add_code(doc: Document, *lines: str):
    """Monospace code block with grey background — one paragraph per line."""
    for line in lines:
        p    = doc.add_paragraph()
        run  = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _set_shading(p, COL_CODE_BG)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
    # Small gap after block
    gap = doc.add_paragraph()
    gap.paragraph_format.space_before = Pt(2)
    gap.paragraph_format.space_after  = Pt(2)


def add_note(doc: Document, text: str, kind: str = "note"):
    """Styled callout box.  kind: 'note' | 'warning' | 'tip'"""
    colours = {
        "note":    (COL_NOTE_BG, "856404", "7B6000"),
        "warning": (COL_WARN_BG, "842029", "CC0000"),
        "tip":     ("D1ECF1",    "0C5460", "0C5460"),
    }
    bg, left, _text_hex = colours.get(kind, colours["note"])
    prefix = {"note": "NOTE  ", "warning": "WARNING  ", "tip": "TIP  "}.get(kind, "NOTE  ")
    p      = doc.add_paragraph()
    run_b  = p.add_run(prefix)
    run_b.bold            = True
    run_b.font.color.rgb  = RGBColor(*bytes.fromhex(_text_hex))
    run_t  = p.add_run(text)
    run_t.font.color.rgb  = RGBColor(*bytes.fromhex(_text_hex))
    _set_para_border(p, bg, left)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    return p


def add_placeholder(doc: Document, label: str,
                    width_in: float = 5.5, height_px: int = 300):
    """Insert a grey placeholder PNG with a caption below."""
    buf = make_placeholder(label, width_px=int(width_in * 96), height_px=height_px)
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(buf, width=Inches(width_in))

    cap = doc.add_paragraph(label)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.size   = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap.paragraph_format.space_after = Pt(8)


def add_bullet(doc: Document, text: str, level: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table_simple(doc: Document, headers: list, rows: list,
                     col_widths: list = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        _set_cell_shading(cell, "1F3964")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        bg  = "FFFFFF" if ri % 2 == 0 else "F5F7FA"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Courier New" if ci > 0 else None
            _set_cell_shading(cell, bg)

    # Column widths
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()   # spacing after table


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Cover ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Installation & User Guide", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in t.runs:
        run.font.color.rgb = COL_HEADING

    sub = doc.add_paragraph(
        "Comparative Evaluation of Vision-Based Finger Tracking for Piano Interaction"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.font.size   = Pt(12)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── 1. System Requirements ───────────────────────────────────────────────
    add_heading(doc, "1.  System Requirements", 1)

    add_body(doc, "Before installing, confirm your system meets the following prerequisites:")
    doc.add_paragraph()

    add_bullet(doc, "Python 3.10  (required — MediaPipe 0.10.x is incompatible with Python 3.11+)")
    add_bullet(doc, "Windows 10 / 11  (primary platform; macOS and Linux work with minor path changes)")
    add_bullet(doc, "Git  (for cloning the repository)")
    add_bullet(doc, "A USB webcam or built-in camera")
    add_bullet(doc, "A MIDI keyboard connected via USB")
    add_bullet(doc, "~2 GB free disk space  (model weights + Python packages)")

    doc.add_paragraph()
    add_placeholder(doc, "Python 3.10 download page at python.org/downloads", height_px=260)

    add_note(doc,
        "Do not use Python 3.11 or higher.  mediapipe==0.10.11 will fail to install "
        "or crash at runtime on anything above 3.10.",
        kind="warning")

    doc.add_paragraph()

    # ── 2. Installation ──────────────────────────────────────────────────────
    add_heading(doc, "2.  Step-by-Step Installation", 1)

    # 2.1 Clone
    add_heading(doc, "2.1  Clone the Repository", 2)
    add_body(doc, "Open a terminal and run:")
    add_code(doc,
        "git clone https://github.com/AayamRegmi/"
        "Comparative-Evaluation-of-Vision-Based-Finger-Tracking-for-Piano-Interaction.git",
        "cd Comparative-Evaluation-of-Vision-Based-Finger-Tracking-for-Piano-Interaction",
    )
    add_placeholder(doc, "Terminal window showing successful git clone output")

    # 2.2 Virtual environment
    add_heading(doc, "2.2  Create a Virtual Environment", 2)
    add_body(doc, "Always use a dedicated virtual environment to avoid package conflicts:")
    add_code(doc, "python -m venv piano_env")
    add_placeholder(doc, "Terminal showing virtual environment creation (piano_env folder created)")

    # 2.3 Activate
    add_heading(doc, "2.3  Activate the Virtual Environment", 2)
    add_body(doc, "Windows (Command Prompt / PowerShell):")
    add_code(doc, r"piano_env\Scripts\activate")
    add_body(doc, "macOS / Linux:")
    add_code(doc, "source piano_env/bin/activate")
    add_body(doc, "Your terminal prompt should show (piano_env) once activated.")
    add_placeholder(doc, "Terminal with (piano_env) prefix visible in the prompt")

    # 2.4 Install requirements
    add_heading(doc, "2.4  Install Python Dependencies", 2)
    add_body(doc,
        "With the virtual environment active, install all required packages from the "
        "requirements file:")
    add_code(doc, "pip install -r requirements.txt")
    add_body(doc,
        "This installs: NumPy, Pandas, OpenCV, MediaPipe, Mido, python-rtmidi, "
        "Matplotlib, Seaborn, SciPy, and tqdm.")
    add_placeholder(doc, "Terminal showing pip install completing with no errors")

    add_note(doc,
        "Windows only — python-rtmidi requires the Microsoft C++ Build Tools to compile.  "
        "If the install fails with a compiler error, download and install "
        "'Build Tools for Visual Studio' from visualstudio.microsoft.com "
        "(select Desktop development with C++), then re-run the pip command.",
        kind="warning")

    add_placeholder(doc, "Visual Studio Build Tools installer — Desktop development with C++ selected",
                    height_px=260)

    doc.add_paragraph()

    # 2.5 Additional packages
    add_heading(doc, "2.5  Install Additional Packages", 2)
    add_body(doc,
        "Three packages used by analysis and results scripts are not included in "
        "requirements.txt and must be installed separately:")
    add_code(doc, "pip install Pillow python-pptx python-docx")
    doc.add_paragraph()

    add_table_simple(doc,
        ["Package", "Version tested", "Used by"],
        [
            ["Pillow",       "12.x",  "live_preview.py — results overlay rendering"],
            ["python-pptx",  "1.x",   "redesign_pptx.py — presentation export"],
            ["python-docx",  "any",   "Documentation helper scripts"],
        ],
        col_widths=[1.4, 1.4, 4.0],
    )

    # 2.6 OpenPose models
    add_heading(doc, "2.6  OpenPose Model Files", 2)
    add_body(doc,
        "OpenPose hand tracking runs through OpenCV's DNN module using pre-trained "
        "Caffe weights.  The model files are already included in the repository "
        "under scripts/ — no separate OpenPose installation is needed:")

    add_bullet(doc, "scripts/openpose_hand.prototxt  — network architecture")
    add_bullet(doc, "scripts/openpose_hand.caffemodel  — pre-trained weights (~55 MB)")

    add_note(doc,
        "If the caffemodel file is missing (e.g. after a shallow clone), run "
        "python download_openpose_model.py from the project root to re-download it.",
        kind="note")
    add_placeholder(doc, "File browser showing scripts/ folder with .prototxt and .caffemodel files",
                    height_px=220)

    doc.add_paragraph()

    # ── 3. Verify ────────────────────────────────────────────────────────────
    add_heading(doc, "3.  Verify the Installation", 1)
    add_body(doc,
        "With the virtual environment active, run the following one-liner to confirm "
        "all core libraries import successfully:")
    add_code(doc,
        'python -c "import cv2, mediapipe, mido, numpy, matplotlib, PIL; '
        'print(\'All core imports OK\')"',
    )
    add_body(doc, "Expected output:")
    add_code(doc, "All core imports OK")
    add_placeholder(doc, "Terminal showing 'All core imports OK' with no errors or warnings")

    doc.add_paragraph()

    # ── 4. MIDI Device Setup ─────────────────────────────────────────────────
    add_heading(doc, "4.  MIDI Device Setup", 1)
    add_body(doc,
        "The system reads MIDI events in real time to log which keys are pressed "
        "during recording and testing sessions.")

    doc.add_paragraph()
    add_bullet(doc, "Connect the MIDI keyboard to the computer via USB before launching any script.")
    add_bullet(doc, "Windows may require a driver from the keyboard manufacturer — check the product page if the keyboard is not detected.")
    add_bullet(doc, "On the setup screen that appears when launching record.py or live_preview.py, the detected MIDI ports are listed.")
    add_bullet(doc, "Use the [ and ] keys to cycle between available ports and select the correct one.")
    add_bullet(doc, "Press a key on the piano to confirm the 'Last note' field updates — this confirms the MIDI connection is working.")

    add_placeholder(doc, "Setup screen showing the MIDI port selection box and 'Last note' live readout")
    add_placeholder(doc, "Windows Device Manager or MIDI driver installation prompt", height_px=240)

    doc.add_paragraph()

    # ── 5. Running the Scripts ───────────────────────────────────────────────
    add_heading(doc, "5.  Running the Scripts", 1)
    add_body(doc,
        "All scripts are run as Python modules from the project root with the "
        "virtual environment active.  The table below summarises each script's purpose:")

    doc.add_paragraph()
    add_table_simple(doc,
        ["Script", "Command", "Purpose"],
        [
            ["key_calibration", "python -m scripts.key_calibration",
             "Align the on-screen keyboard overlay to the camera view and save calibration"],
            ["live_preview",    "python -m scripts.live_preview",
             "Real-time MPJPE accuracy test — no data written to disk"],
            ["record",          "python -m scripts.record",
             "Record a labelled session: video, MIDI, and session metadata"],
            ["analyse",         "python -m scripts.analyse",
             "Offline MPJPE analysis of recorded sessions via GUI or CLI"],
            ["plot_results",    "python -m scripts.plot_results",
             "Generate all charts and participant reports from analysed data"],
        ],
        col_widths=[1.5, 2.5, 3.3],
    )

    # 5.1
    add_heading(doc, "5.1  Key Calibration  (key_calibration.py)", 2)
    add_body(doc,
        "Run this first, or whenever the camera position changes.  Drag the "
        "keyboard overlay to match the physical piano keys in the camera view, "
        "then press ENTER to save.")
    add_placeholder(doc, "key_calibration.py running — keyboard overlay aligned to physical piano keys")

    add_body(doc, "Key controls:")
    add_table_simple(doc,
        ["Key", "Action"],
        [
            ["Drag interior",    "Move the whole overlay"],
            ["Drag edges",       "Resize key width or height"],
            ["Drag corner ◯",    "Warp perspective to correct camera angle"],
            ["Scroll wheel",     "Add or remove a key on the right"],
            ["[ / ]",            "Shift the starting MIDI note"],
            ["C",                "Toggle centre visualisation (line / dot)"],
            ["F",                "Toggle 180° flip to match record.py"],
            ["H",                "Toggle help overlay"],
            ["ENTER or S",       "Save calibration"],
            ["ESC",              "Quit without saving"],
        ],
        col_widths=[1.8, 4.5],
    )

    # 5.2
    add_heading(doc, "5.2  Live MPJPE Test  (live_preview.py)", 2)
    add_body(doc,
        "Opens a setup screen to select the MIDI port and camera.  Once running, "
        "play the piano and MPJPE accuracy is shown live in the top-right corner.  "
        "Press ESC to stop and view the final results panel.")
    add_placeholder(doc, "live_preview.py main view — hand skeleton, key overlay, MPJPE readout top-right")
    add_placeholder(doc, "live_preview.py results panel showing MPJPE, accuracy %, and per-finger breakdown",
                    height_px=360)

    add_body(doc, "Key controls during test:")
    add_table_simple(doc,
        ["Key", "Action"],
        [
            ["M",   "Toggle key mask overlay on/off"],
            ["N",   "Toggle mask control panel"],
            ["C",   "Toggle centre visualisation (line / dot)"],
            ["S",   "Toggle FPS / latency stats overlay"],
            ["V",   "Reset perspective warp to rectangle"],
            ["ESC", "Stop test and show results"],
        ],
        col_widths=[1.0, 5.3],
    )

    # 5.3
    add_heading(doc, "5.3  Recording Session  (record.py)", 2)
    add_body(doc,
        "Records a full participant session.  The setup screen collects participant "
        "ID, Fitzpatrick skin type, ambient lux, and hand size before recording "
        "begins.  Video, MIDI, and metadata are saved to data/raw/p###/.")
    add_placeholder(doc, "record.py setup screen — participant ID, Fitzpatrick type, lux, and hand-size fields")
    add_placeholder(doc, "record.py recording view — participant playing piano with overlay active")

    # 5.4
    add_heading(doc, "5.4  Offline Analysis  (analyse.py)", 2)
    add_body(doc,
        "Processes recorded sessions to compute MPJPE metrics for MediaPipe and/or "
        "OpenPose.  Can be run via the GUI (double-click sessions to process) or "
        "from the command line:")
    add_code(doc,
        "python -m scripts.analyse                       # open GUI",
        "python -m scripts.analyse data/raw/p001/        # process one session",
        "python -m scripts.analyse data/raw/p001/ --model openpose",
    )
    add_placeholder(doc, "analyse.py Tkinter GUI showing session list with MediaPipe and OpenPose columns")

    # 5.5
    add_heading(doc, "5.5  Plot Results  (plot_results.py)", 2)
    add_body(doc,
        "Reads all results JSON files under data/processed/ and generates group-level "
        "charts and per-participant PDFs in data/plots/:")
    add_code(doc, "python -m scripts.plot_results")
    add_placeholder(doc, "data/plots/ folder — generated PNG charts visible in file browser")

    doc.add_paragraph()

    # ── 6. Data Folder Structure ─────────────────────────────────────────────
    add_heading(doc, "6.  Data Folder Structure", 1)
    add_body(doc,
        "The following folder structure is created automatically as sessions are "
        "recorded and processed:")
    add_code(doc,
        "data/",
        "├── calibration/",
        "│   └── key_centers.json          ← saved by key_calibration.py",
        "├── raw/",
        "│   └── p001/",
        "│       ├── p001.mp4              ← raw video (no overlays)",
        "│       ├── p001_frames.csv       ← frame timestamps",
        "│       ├── p001_midi.jsonl       ← MIDI events",
        "│       ├── p001_midi.mid         ← standard MIDI file",
        "│       └── p001_session.json     ← session metadata",
        "├── processed/",
        "│   └── p001/",
        "│       ├── p001_mediapipe_results.json",
        "│       └── p001_openpose_results.json",
        "└── plots/",
        "    ├── 01_model_comparison_mpjpe.png",
        "    └── ...                        ← all group charts",
    )

    doc.add_paragraph()

    # ── 7. Troubleshooting ───────────────────────────────────────────────────
    add_heading(doc, "7.  Troubleshooting", 1)

    add_table_simple(doc,
        ["Problem", "Likely cause", "Fix"],
        [
            ["pip install fails on python-rtmidi",
             "Missing C++ compiler",
             "Install Visual Studio Build Tools (Desktop development with C++)"],
            ["ModuleNotFoundError: mediapipe",
             "Wrong Python version or venv not active",
             "Confirm python --version is 3.10.x and (piano_env) prefix is shown"],
            ["No MIDI ports detected",
             "Keyboard not connected or driver missing",
             "Connect keyboard before launching; install manufacturer USB-MIDI driver"],
            ["Camera shows black / wrong camera",
             "Wrong camera index in config.py",
             "Edit CAMERA_INDEX in scripts/config.py (try 0, 1, 2…)"],
            ["caffemodel missing error",
             "Large file not pulled from Git LFS",
             "Run python download_openpose_model.py to re-download"],
            ["All core imports OK but MediaPipe crashes",
             "Python 3.11+ installed",
             "Switch to Python 3.10 and recreate the virtual environment"],
        ],
        col_widths=[2.0, 2.2, 3.1],
    )

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
