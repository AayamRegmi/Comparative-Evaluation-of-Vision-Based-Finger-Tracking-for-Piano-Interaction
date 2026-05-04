"""
Builds Documentation/User Guidelines.docx from scratch.
Run with:  piano_env\Scripts\python.exe build_user_guide.py
"""

from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from PIL import Image, ImageDraw, ImageFont

OUT_PATH = Path("Documentation/User Guidelines.docx")

# ---------------------------------------------------------------------------
# Colour palette
# ---------------------------------------------------------------------------

COL_HEADING   = RGBColor(0x1F, 0x39, 0x64)
COL_CODE_BG   = "F2F2F2"
COL_NOTE_BG   = "FFF3CD"
COL_INFO_BG   = "D1ECF1"
COL_PLACEHOLDER = (180, 185, 195)
COL_PH_BORDER   = (130, 135, 145)
COL_PH_TEXT     = (80, 85, 95)

# ---------------------------------------------------------------------------
# Placeholder image generator
# ---------------------------------------------------------------------------

def make_placeholder(label: str, width_px: int = 580, height_px: int = 300) -> BytesIO:
    img  = Image.new("RGB", (width_px, height_px), COL_PLACEHOLDER)
    draw = ImageDraw.Draw(img)
    bw   = 3
    draw.rectangle([bw, bw, width_px - bw - 1, height_px - bw - 1],
                   outline=COL_PH_BORDER, width=bw)

    cx, cy = width_px // 2, height_px // 2 - 24
    iw, ih = 60, 44
    draw.rectangle([cx - iw, cy - ih, cx + iw, cy + ih], outline=COL_PH_TEXT, width=2)
    draw.polygon([(cx - iw + 8, cy + ih - 2), (cx - 8, cy - ih + 16),
                  (cx + 26, cy + ih - 2)], outline=COL_PH_TEXT)
    draw.polygon([(cx + 8, cy + ih - 2), (cx + iw - 6, cy - ih + 24),
                  (cx + iw - 2, cy + ih - 2)], outline=COL_PH_TEXT)
    draw.ellipse([cx + 28, cy - ih + 4, cx + 50, cy - ih + 22], outline=COL_PH_TEXT)

    try:
        font_main  = ImageFont.truetype("arialbd.ttf", 17)
        font_small = ImageFont.truetype("arial.ttf",   13)
    except OSError:
        font_main = font_small = ImageFont.load_default()

    main_text = "[ IMAGE PLACEHOLDER ]"
    bb  = draw.textbbox((0, 0), main_text, font=font_main)
    draw.text(((width_px - (bb[2]-bb[0])) // 2, cy + ih + 10),
              main_text, font=font_main, fill=COL_PH_TEXT)
    bb2 = draw.textbbox((0, 0), label, font=font_small)
    draw.text(((width_px - (bb2[2]-bb2[0])) // 2, cy + ih + 36),
              label, font=font_small, fill=COL_PH_TEXT)

    buf = BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# XML helpers
# ---------------------------------------------------------------------------

def _shd(paragraph, fill_hex: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _cell_shd(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _left_bar(paragraph, fill_hex: str, bar_color: str = "1F3964"):
    _shd(paragraph, fill_hex)
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), bar_color)
    pBdr.append(left)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = COL_HEADING
    return p


def add_body(doc, text):
    return doc.add_paragraph(text)


def add_code(doc, *lines):
    for line in lines:
        p   = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name  = "Courier New"
        run.font.size  = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        _shd(p, COL_CODE_BG)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        p.paragraph_format.left_indent  = Cm(0.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)


def add_note(doc, text, kind="note"):
    colours = {
        "note":  (COL_NOTE_BG, "856404", "7B6000"),
        "info":  (COL_INFO_BG, "0C5460", "0C5460"),
    }
    bg, left, txt_hex = colours.get(kind, colours["note"])
    prefix = {"note": "NOTE  ", "info": "INFO  "}.get(kind, "NOTE  ")
    p      = doc.add_paragraph()
    rb     = p.add_run(prefix)
    rb.bold = True
    rb.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    rt = p.add_run(text)
    rt.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    _left_bar(p, bg, left)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_placeholder(doc, label, width_in=5.5, height_px=280):
    buf = make_placeholder(label, width_px=int(width_in * 96), height_px=height_px)
    p   = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(buf, width=Inches(width_in))
    cap = doc.add_paragraph(label)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.font.size   = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    cap.paragraph_format.space_after = Pt(8)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def add_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    # header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _cell_shd(c, "1F3964")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        bg  = "FFFFFF" if ri % 2 == 0 else "F5F7FA"
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            if ci > 0:
                r.font.name = "Courier New"
            _cell_shd(c, bg)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Script entry card helper
# ---------------------------------------------------------------------------

def script_card(doc, script_name, purpose, depends_on, used_by=None):
    """Renders the standard Purpose / Depends on / Used by block for a script."""
    p = doc.add_paragraph()
    r = p.add_run("Purpose:  ")
    r.bold = True
    p.add_run(purpose)

    p2 = doc.add_paragraph()
    r2 = p2.add_run("Depends on:  ")
    r2.bold = True
    p2.add_run(depends_on)

    if used_by:
        p3 = doc.add_paragraph()
        r3 = p3.add_run("Used by:  ")
        r3.bold = True
        p3.add_run(used_by)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Script Reference  &  User Guide", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.color.rgb = COL_HEADING

    sub = doc.add_paragraph(
        "Comparative Evaluation of Vision-Based Finger Tracking for Piano Interaction"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size   = Pt(12)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── 1. Pipeline Overview ─────────────────────────────────────────────────
    add_heading(doc, "1.  Pipeline Overview", 1)
    add_body(doc,
        "The system is structured as a sequential pipeline.  Each script hands off "
        "its output to the next stage.  live_preview.py is a standalone tool that "
        "does not persist any data.")
    doc.add_paragraph()

    add_code(doc,
        "key_calibration.py          run once per camera/tripod position",
        "        |",
        "record.py                   capture video + MIDI per participant",
        "        |",
        "recalibrate.py              (optional) fix calibration on saved video",
        "        |",
        "analyse.py                  compute MPJPE results per session",
        "        |",
        "plot_results.py             generate group charts and participant reports",
        "",
        "live_preview.py             standalone live test  (no data written to disk)",
        "data_summary.py             audit data/raw/ session counts at any time",
    )

    add_placeholder(doc,
        "Pipeline flow diagram — key_calibration to record to analyse to plot_results",
        height_px=260)

    # ── 2. Configuration ─────────────────────────────────────────────────────
    add_heading(doc, "2.  config.py", 1)

    script_card(doc,
        "config.py",
        "Central store for every tunable constant in the project.  All other "
        "scripts import this instead of hard-coding values.",
        "nothing  (no internal imports)",
        used_by="all scripts",
    )

    add_heading(doc, "Key constants", 2)
    add_table(doc,
        ["Constant", "Description"],
        [
            ["CAP_WIDTH / HEIGHT / FPS",
             "Camera capture resolution and target frame rate"],
            ["RESIZE_WIDTH / HEIGHT",
             "Resolution MediaPipe runs inference at (smaller = faster)"],
            ["MODEL_COMPLEXITY",
             "MediaPipe model complexity  (0 = lite, 1 = full)"],
            ["MIN_DETECTION_CONFIDENCE",
             "Minimum score to accept a hand detection"],
            ["MIN_TRACKING_CONFIDENCE",
             "Minimum score to continue tracking a detected hand"],
            ["LUX_DIM_THRESHOLD",
             "Below this value (lux) the session is labelled Dim"],
            ["LUX_BRIGHT_THRESHOLD",
             "At or above this value the session is labelled Bright"],
            ["HAND_SIZE_MIN / MAX_CM",
             "Validation range for the manually entered hand length"],
            ["HAND_SIZE_SMALL_MAX_CM",
             "Hand lengths below this are categorised as Small"],
            ["HAND_SIZE_LARGE_MIN_CM",
             "Hand lengths at or above this are categorised as Large"],
            ["ACCURACY_THRESHOLD_RATIO",
             "Fraction of white key width used as the accuracy pass threshold"],
            ["MASK_*",
             "Fill and border colours for the keyboard overlay drawn on-screen"],
        ],
        col_widths=[2.4, 4.9],
    )

    # ── 3. Support scripts ───────────────────────────────────────────────────
    add_heading(doc, "3.  Support Scripts", 1)
    add_body(doc,
        "These scripts are imported by the main pipeline scripts.  "
        "They are not run directly.")
    doc.add_paragraph()

    # camera_setup
    add_heading(doc, "3.1  camera_setup.py", 2)
    script_card(doc,
        "camera_setup",
        "Opens the camera with the correct backend for the current OS and applies "
        "the capture resolution and FPS from config.  Returns a ready VideoCapture object.",
        "config,  cv2",
        used_by="record.py,  key_calibration.py,  live_preview.py",
    )
    add_body(doc, "Key function:")
    add_bullet(doc,
        "init_camera(config)  —  opens the camera using CAP_DSHOW on Windows "
        "(required to avoid DirectShow double-open conflicts), sets width/height/FPS "
        "and buffer size, returns the VideoCapture handle.")
    doc.add_paragraph()

    # lux_calculator
    add_heading(doc, "3.2  lux_calculator.py", 2)
    script_card(doc,
        "lux_calculator",
        "Converts a numeric lux value into a human-readable lighting category label.",
        "config",
        used_by="record.py,  live_preview.py",
    )
    add_body(doc, "Key function:")
    add_bullet(doc,
        'lux_to_label(lux)  ->  "Dim" / "Indoor" / "Bright"  '
        "(thresholds are set in config.py)")
    doc.add_paragraph()

    # fitzpatrick_detector
    add_heading(doc, "3.3  fitzpatrick_detector.py", 2)
    script_card(doc,
        "fitzpatrick_detector",
        "Estimates the participant's Fitzpatrick skin type (I-VI) from a live video "
        "frame using the ITA (Individual Typology Angle) colorimetric method.  "
        "Samples CIELab pixels from the palm convex hull (landmarks 0,1,5,9,13,17).  "
        "Runs once per session at setup time.",
        "config,  cv2,  numpy",
        used_by="record.py",
    )
    add_body(doc, "Key functions:")
    add_bullet(doc,
        "detect_skin_type(frame, hand_landmarks, fw, fh)  "
        "->  (fitzpatrick_type, label, ita_angle)  or  (None, None, None) if palm not visible")
    add_bullet(doc,
        "ita_to_fitzpatrick(ita)  ->  (type_int, label_str)  "
        "— maps ITA angle to the six Fitzpatrick categories using Chardon (1991) thresholds")
    doc.add_paragraph()

    # stats_collector
    add_heading(doc, "3.4  stats_collector.py", 2)
    script_card(doc,
        "stats_collector",
        "Rolling statistics tracker for the recording/preview loop.  Computes FPS, "
        "MediaPipe inference time, and total frame latency over a configurable window.",
        "numpy,  collections.deque",
        used_by="record.py,  live_preview.py",
    )
    add_body(doc, "Key class:  StatsCollector(warmup_frames, stat_frames)")
    add_table(doc,
        ["Method / property", "Description"],
        [
            ["update(loop_start, inference_ms, latency_ms)",
             "Call once per frame to push new measurements"],
            ["fps",
             "Rolling average frames per second"],
            ["inference_ms",
             "Rolling average MediaPipe inference time (ms)"],
            ["latency_ms",
             "Rolling average total frame latency (ms)"],
            ["warmup_done",
             "True once warmup_frames have elapsed"],
            ["print_final_stats(w, h, complexity)",
             "Prints a summary table on session exit"],
        ],
        col_widths=[3.0, 4.3],
    )

    # midi_recorder
    add_heading(doc, "3.5  midi_recorder.py", 2)
    script_card(doc,
        "midi_recorder",
        "Threaded MIDI input capture with time.perf_counter() timestamps.  "
        "All timestamps are relative to a shared t0 set at recording start "
        "(SPACE keypress), keeping video and MIDI in sync.",
        "mido,  python-rtmidi,  json,  threading,  pathlib",
        used_by="record.py,  live_preview.py",
    )
    add_body(doc, "Key classes:")
    add_table(doc,
        ["Class / function", "Description"],
        [
            ["MidiRecorder.start(t0)",
             "Begins background daemon thread polling the MIDI port every 1 ms.  "
             "Skips realtime messages (clock, active sensing)."],
            ["MidiRecorder.stop()",
             "Joins the background thread."],
            ["MidiRecorder.save(out_dir, pid_str)",
             "Writes p###_midi.jsonl (one JSON line per event) and p###_midi.mid "
             "(standard MIDI file, type 0, 120 BPM)."],
            ["FrameLogger.log()",
             "Call once per writer.write(frame).  Records (frame_index, time_s) "
             "relative to t0."],
            ["FrameLogger.save(out_dir, pid_str)",
             "Writes p###_frames.csv."],
            ["list_midi_input_ports()",
             "Returns list of detected MIDI port name strings."],
        ],
        col_widths=[2.5, 4.8],
    )

    # model_manager
    add_heading(doc, "3.6  model_manager.py", 2)
    script_card(doc,
        "model_manager",
        "Abstract interface and concrete wrappers for switching between hand-tracking "
        "models at runtime.  Defines a common API (load, infer, draw, close) so "
        "analyse.py can run either model without branching logic.",
        "config,  cv2,  numpy,  mediapipe",
        used_by="live_preview.py,  analyse.py",
    )
    add_body(doc, "Concrete model classes:")
    add_bullet(doc,
        "MediaPipeHandModel  —  wraps mp.solutions.hands.  "
        "infer() returns the standard MediaPipe result object.")
    add_bullet(doc,
        "OpenPoseHandModel  —  loads pre-trained Caffe weights via cv2.dnn.  "
        "infer() extracts 21 keypoints from heatmap outputs.  "
        "Operates on a keyboard-area crop rather than the full frame.")
    doc.add_paragraph()

    # ── 4. Key Calibration ───────────────────────────────────────────────────
    add_heading(doc, "4.  key_calibration.py", 1)

    script_card(doc,
        "key_calibration",
        "Interactive GUI for aligning a virtual piano keyboard overlay over the "
        "camera view.  The overlay geometry (position, key width/height, perspective "
        "warp) is saved to data/calibration/key_centers.json and loaded by all "
        "downstream scripts.",
        "config,  camera_setup,  cv2,  numpy,  json,  math",
        used_by="record.py,  live_preview.py,  analyse.py,  recalibrate.py",
    )

    add_body(doc, "Key class:  KeyMask")
    add_table(doc,
        ["Method", "Description"],
        [
            ["draw(frame, center_viz='line')",
             "Blends the perspective-warped key overlay onto a frame in-place.  "
             "center_viz='line' draws a full-depth vertical line per key; "
             "'dot' draws a single point at the calculated centre."],
            ["on_mouse(event, x, y, flags)",
             "Handles drag-to-move, edge-drag-to-resize, "
             "corner-drag-to-warp, scroll-to-add/remove keys."],
            ["save(cal_dir, fw, fh)",
             "Writes key_centers.json with layout parameters and "
             "pre-computed per-key centre coordinates."],
            ["KeyMask.load(path)",
             "Reconstructs a KeyMask from a saved JSON file."],
            ["KeyMask.default(fw, fh)",
             "Returns a default 29-key layout sized to ~3/4 frame width."],
        ],
        col_widths=[2.6, 4.7],
    )

    add_body(doc, "Exported utility function:")
    add_bullet(doc,
        "build_key_layout(start_midi, num_white, ox, oy, wkw, wkh)  "
        "—  re-derives per-key centre coordinates from saved parameters; "
        "used by analyse.py and recalibrate.py.")
    doc.add_paragraph()

    add_heading(doc, "Controls", 2)
    add_table(doc,
        ["Key / Action", "Effect"],
        [
            ["Drag interior",         "Move the entire keyboard overlay"],
            ["Drag left / right edge", "Resize key width (opposite edge stays fixed)"],
            ["Drag bottom edge",      "Resize key height"],
            ["Drag corner circle",    "Warp perspective to correct camera angle"],
            ["Scroll wheel",          "Add or remove one key on the right end"],
            ["[ / ]",                 "Shift the starting MIDI note down / up"],
            ["- / =",                 "Fine key width -1 / +1 px"],
            ["C",                     "Toggle centre visualisation  (line / dot)"],
            ["F",                     "Toggle 180 degree flip to match record.py orientation"],
            ["H",                     "Toggle on-screen help"],
            ["V",                     "Reset perspective warp to rectangle"],
            ["ENTER or S",            "Save calibration to key_centers.json"],
            ["ESC",                   "Quit without saving"],
        ],
        col_widths=[2.0, 5.3],
    )

    add_placeholder(doc,
        "key_calibration.py — keyboard overlay aligned to physical piano keys on camera feed")

    add_note(doc,
        "Run key_calibration.py first, or any time the camera or tripod position "
        "changes.  All downstream scripts rely on the saved key_centers.json.",
        kind="info")

    doc.add_paragraph()

    # ── 5. record.py ────────────────────────────────────────────────────────
    add_heading(doc, "5.  record.py", 1)

    script_card(doc,
        "record",
        "Main data collection script.  Displays a live MediaPipe hand-tracking "
        "preview with the key overlay active.  Before each participant, a setup "
        "screen collects participant ID, Fitzpatrick skin type, ambient lux, and "
        "hand size.  Pressing SPACE starts and stops recording.",
        "config,  camera_setup,  fitzpatrick_detector,  lux_calculator,  "
        "stats_collector,  midi_recorder,  key_calibration,  mediapipe,  "
        "cv2,  numpy,  mido",
        used_by="(main script — not imported by others)",
    )

    add_heading(doc, "Output files per session  (data/raw/p###/)", 2)
    add_table(doc,
        ["File", "Contents"],
        [
            ["p###.mp4",
             "Raw video with no overlays.  180 degree flip baked in if flip_y = True."],
            ["p###_frames.csv",
             "frame_index, time_s — one row per recorded frame, timestamps "
             "relative to recording start."],
            ["p###_midi.jsonl",
             "One JSON line per MIDI event with time_s offset from recording start."],
            ["p###_midi.mid",
             "Standard MIDI file (type 0, 120 BPM, 480 ticks/beat)."],
            ["p###_session.json",
             "Session metadata: lux, hand size, Fitzpatrick type, flip state, "
             "MIDI port name, and references to all output file names."],
            ["p###_key_centers.json",
             "Snapshot of the active calibration at the moment recording started."],
        ],
        col_widths=[2.0, 5.3],
    )

    add_heading(doc, "Key controls", 2)
    add_table(doc,
        ["Key", "Action"],
        [
            ["SPACE",  "Start recording  /  stop recording"],
            ["N",      "Finish session and go to setup screen for next participant"],
            ["M",      "Toggle key overlay visibility"],
            ["P",      "Toggle mask control panel (move/resize overlay with buttons)"],
            ["C",      "Toggle centre visualisation  (line / dot)"],
            ["S",      "Toggle FPS / inference / latency stats overlay"],
            ["R",      "Re-run Fitzpatrick skin-type auto-detection"],
            ["F",      "Toggle 180 degree frame flip"],
            ["V",      "Reset mask perspective warp to rectangle"],
            ["ESC",    "Quit"],
        ],
        col_widths=[1.0, 6.3],
    )

    add_placeholder(doc,
        "record.py setup screen — participant ID, Fitzpatrick type, lux, and hand-size input fields")
    add_placeholder(doc,
        "record.py recording view — hand skeleton and key overlay active, REC indicator visible")

    doc.add_paragraph()

    # ── 6. live_preview.py ──────────────────────────────────────────────────
    add_heading(doc, "6.  live_preview.py", 1)

    script_card(doc,
        "live_preview",
        "Live hand-tracking preview and real-time MPJPE accuracy test.  "
        "Mirrors record.py's display but writes nothing to disk.  "
        "Shows live MPJPE, accuracy percentage, and detection rate in the top-right "
        "corner.  On exit (ESC), displays a full results panel with per-finger breakdown.",
        "config,  camera_setup,  key_calibration,  midi_recorder,  "
        "stats_collector,  mediapipe,  PIL,  cv2,  numpy,  mido",
        used_by="(standalone tool — not imported by others)",
    )

    add_heading(doc, "MPJPE computation", 2)
    add_body(doc,
        "For every MIDI note_on event (velocity > 0) while at least one hand is "
        "visible, the script classifies the outcome into one of three categories:")
    add_table(doc,
        ["Outcome", "Meaning", "Effect on MPJPE"],
        [
            ["matched",
             "A fingertip landmark was inside the key polygon",
             "Horizontal error |tip_x - key_cx| added to MPJPE"],
            ["detection_fail",
             "Hands visible but no fingertip tip inside the polygon",
             "Counted separately — NOT added to MPJPE"],
            ["missed",
             "No hand landmarks detected at all",
             "Counted separately — NOT added to MPJPE"],
        ],
        col_widths=[1.5, 3.0, 2.8],
    )

    add_note(doc,
        "MPJPE uses horizontal-only error  (|tip_x - key_cx|)  to remove "
        "depth / finger-length bias from the measurement.  "
        "Any fingertip Y position within the key height is considered correct.",
        kind="info")

    add_heading(doc, "Key controls", 2)
    add_table(doc,
        ["Key", "Action"],
        [
            ["M",   "Toggle key mask overlay"],
            ["N",   "Toggle mask control panel"],
            ["C",   "Toggle centre visualisation  (line / dot)"],
            ["S",   "Toggle FPS / inference / latency stats overlay"],
            ["V",   "Reset perspective warp to rectangle"],
            ["ESC", "Stop test and show results panel"],
        ],
        col_widths=[1.0, 6.3],
    )

    add_placeholder(doc,
        "live_preview.py — hand skeleton, key overlay, and live MPJPE readout top-right corner")
    add_placeholder(doc,
        "live_preview.py results panel — MPJPE, accuracy %, detection rate, and per-finger table",
        height_px=340)

    doc.add_paragraph()

    # ── 7. recalibrate.py ───────────────────────────────────────────────────
    add_heading(doc, "7.  recalibrate.py", 1)

    script_card(doc,
        "recalibrate",
        "GUI tool for aligning the key overlay against a pre-recorded session video.  "
        "Useful when the session calibration file was lost or was inaccurate.  "
        "Saves a per-session p###_key_centers.json that analyse.py will prefer "
        "over the global calibration.",
        "key_calibration.KeyMask,  cv2,  numpy,  json,  csv",
        used_by="(standalone tool — run manually before re-analysing a session)",
    )

    add_heading(doc, "Workflow", 2)
    add_body(doc,
        "1.  Run  python -m scripts.recalibrate  to open the participant picker.")
    add_body(doc,
        "2.  A scrollable list of all data/raw/p###/ folders is shown.  "
        "Green dot = already has a per-session calibration.  "
        "Blue dot = will use the global calibration.")
    add_body(doc,
        "3.  Click a participant to open the calibration editor, which plays back "
        "that session's video.")
    add_body(doc,
        "4.  Align the overlay to match the piano in the video.  "
        "The NOW: C5 indicator shows which MIDI note was played at the current frame.")
    add_body(doc,
        "5.  Press ENTER or S to save p###_key_centers.json and return to the picker.")
    doc.add_paragraph()

    add_placeholder(doc,
        "recalibrate.py participant picker — scrollable list with green/blue status dots")
    add_placeholder(doc,
        "recalibrate.py calibration editor — video playback with overlay and MIDI note indicator")

    doc.add_paragraph()

    # ── 8. analyse.py ───────────────────────────────────────────────────────
    add_heading(doc, "8.  analyse.py", 1)

    script_card(doc,
        "analyse",
        "Offline per-session MPJPE computation.  For each note_on event in the "
        "MIDI log, it seeks to the nearest video frame, runs the chosen model, "
        "and records the outcome.  Supports MediaPipe and OpenPose via model_manager.  "
        "Can be run from the command line or via a Tkinter GUI session picker.",
        "config,  key_calibration.KeyMask,  model_manager,  mediapipe,  cv2,  numpy",
        used_by="(main analysis script — outputs consumed by plot_results.py)",
    )

    add_heading(doc, "Calibration priority", 2)
    add_body(doc,
        "analyse.py looks for a per-session calibration file first "
        "(data/raw/p###/p###_key_centers.json), then falls back to the global "
        "calibration (data/calibration/key_centers.json).  "
        "Use recalibrate.py to create per-session files where needed.")
    doc.add_paragraph()

    add_heading(doc, "Hand classification", 2)
    add_bullet(doc,
        "MediaPipe:  uses multi_handedness label combined with thumb-MCP-to-wrist "
        "geometry to determine physical L / R, accounting for mirrored camera setups.")
    add_bullet(doc,
        "OpenPose:  geometry is unreliable on scene-level crops; falls back to "
        "keyboard-position split  (left of keyboard midpoint = L).")
    doc.add_paragraph()

    add_heading(doc, "Output files  (data/processed/p###/)", 2)
    add_table(doc,
        ["File", "Contents"],
        [
            ["p###_mediapipe_results.json",
             "Per-session MPJPE results: detection_rate_pct, per_hand L/R with "
             "matched, detection_fail, mpjpe_px, accuracy_pct, and per-finger arrays."],
            ["p###_openpose_results.json",
             "Same structure for OpenPose (hand split uses keyboard-position fallback)."],
        ],
        col_widths=[2.8, 4.5],
    )

    add_heading(doc, "Usage", 2)
    add_code(doc,
        "python -m scripts.analyse                       # open Tkinter GUI",
        "python -m scripts.analyse data/raw/p001/        # process one session (both models)",
        "python -m scripts.analyse data/raw/p001/ --model mediapipe",
        "python -m scripts.analyse data/raw/p001/ --model openpose",
    )

    add_placeholder(doc,
        "analyse.py Tkinter GUI — session list with MediaPipe and OpenPose status columns")

    doc.add_paragraph()

    # ── 9. plot_results.py ──────────────────────────────────────────────────
    add_heading(doc, "9.  plot_results.py", 1)

    script_card(doc,
        "plot_results",
        "Reads all processed result JSON files under data/processed/ and generates "
        "publication-ready group-level charts plus individual participant reports "
        "in data/plots/.",
        "matplotlib,  numpy,  json,  pathlib",
        used_by="(terminal script — outputs are PNG files for the report)",
    )

    add_heading(doc, "Generated charts  (data/plots/)", 2)
    add_table(doc,
        ["File", "Description"],
        [
            ["01_model_comparison_mpjpe.png",
             "Bar chart: MediaPipe vs OpenPose MPJPE per participant"],
            ["02_per_finger_mpjpe.png",
             "Mean per-finger MPJPE: MediaPipe L, MediaPipe R, OpenPose combined"],
            ["02_per_finger_mpjpe_mediapipe.png",
             "MediaPipe-only per-finger MPJPE: L and R panels"],
            ["03_detection_breakdown.png",
             "Stacked bar: matched / detection-fail / missed proportions per model"],
            ["04_mpjpe_by_lux.png",
             "MPJPE grouped by lighting condition (Dim / Indoor / Bright) with SD error bars"],
            ["05_mpjpe_by_fitzpatrick.png",
             "MPJPE grouped by Fitzpatrick skin type with SD error bars"],
            ["06_mpjpe_vs_handsize.png",
             "Scatter: MPJPE vs hand size (cm) with regression lines per model"],
            ["07_finger_distribution_mediapipe.png",
             "Box plots of per-finger MPJPE distribution across all sessions (MediaPipe)"],
            ["07_finger_distribution_openpose.png",
             "Same for OpenPose (combined L+R)"],
            ["08_heatmap_mediapipe.png",
             "Heatmap: participants x fingers, 10 columns (L/R split)"],
            ["08_heatmap_openpose.png",
             "Heatmap: participants x fingers, 5 columns (combined)"],
            ["reports/p###_report.png",
             "Individual participant report (see below)"],
        ],
        col_widths=[3.0, 4.3],
    )

    add_heading(doc, "Participant report layout", 2)
    add_table(doc,
        ["Panel position", "Contents"],
        [
            ["Top left",    "Session info: PID, Fitzpatrick type, lighting, hand size, notes played"],
            ["Top centre",  "Detection outcome bar chart  (Matched / Det.Fail / Missed)"],
            ["Top right",   "Overall results: MPJPE, accuracy %, detection rate for both models"],
            ["Bottom left", "Per-hand bar chart (MediaPipe) coloured by accuracy tier"],
            ["Bottom centre", "Per-hand bar chart (OpenPose) coloured by accuracy tier"],
            ["Bottom right", "Highlights: best/worst finger, detection rate rating, overall accuracy rating"],
        ],
        col_widths=[1.8, 5.5],
    )

    add_heading(doc, "Usage", 2)
    add_code(doc,
        "python -m scripts.plot_results                       # all sessions",
        "python -m scripts.plot_results --pid p001 p003       # specific sessions",
        "python -m scripts.plot_results --out figures/        # custom output folder",
    )

    add_placeholder(doc,
        "Example generated group chart — model comparison MPJPE bar chart")
    add_placeholder(doc,
        "Example participant report — p001_report.png with all 6 panels", height_px=340)

    doc.add_paragraph()

    # ── 10. data_summary.py ─────────────────────────────────────────────────
    add_heading(doc, "10.  data_summary.py", 1)

    script_card(doc,
        "data_summary",
        "Quick audit of how much data has been collected across all sessions.  "
        "Prints a table with per-session frame counts, MIDI event counts, and "
        "matched note statistics.  Does not depend on any internal scripts.",
        "json,  csv,  pathlib",
        used_by="(standalone audit tool — no imports of project scripts)",
    )

    add_code(doc, "python -m scripts.data_summary")

    add_table(doc,
        ["Column", "Description"],
        [
            ["Frames",       "Total video frames recorded in the session"],
            ["MIDI evts",    "Total MIDI messages of all types"],
            ["note_on",      "Key-press events with velocity > 0"],
            ["MP matched",   "Notes where a MediaPipe fingertip landed inside the key polygon"],
            ["OP matched",   "Notes where an OpenPose fingertip landed inside the key polygon"],
        ],
        col_widths=[1.5, 5.8],
    )

    doc.add_paragraph()

    # ── 11. Dependency Graph ─────────────────────────────────────────────────
    add_heading(doc, "11.  Dependency Graph", 1)
    add_body(doc,
        "Arrows show import relationships.  scripts at the top are imported by "
        "everything below them.")
    doc.add_paragraph()

    add_code(doc,
        "config.py",
        "  camera_setup.py",
        "  lux_calculator.py",
        "  fitzpatrick_detector.py",
        "  stats_collector.py",
        "  midi_recorder.py",
        "  key_calibration.py",
        "  |   recalibrate.py",
        "  model_manager.py",
        "  |   live_preview.py  ............  (no data written to disk)",
        "  |   analyse.py",
        "  |       plot_results.py",
        "  record.py  (imports all of the above)",
        "",
        "  data_summary.py  (reads data/raw/ directly, imports no project scripts)",
    )

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
