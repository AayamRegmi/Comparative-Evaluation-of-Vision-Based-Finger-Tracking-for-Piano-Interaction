"""
Builds Documentation/Requirement Catalogue.docx
Run with:  piano_env\Scripts\python.exe build_requirement_catalogue.py
"""

from pathlib import Path
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = Path("Documentation/Requirement Catalogue.docx")

COL_HEADING  = RGBColor(0x1F, 0x39, 0x64)
COL_MUST     = "1F3964"   # dark navy   header
COL_SHOULD   = "2E6B9E"   # mid blue
COL_COULD    = "4A9CC7"   # light blue
COL_WONT     = "7F7F7F"   # grey
COL_DONE     = "D6EAD6"   # pale green  row bg
COL_PARTIAL  = "FFF3CD"   # amber
COL_NOT_DONE = "F8D7DA"   # pale red
COL_EXCLUDED = "F2F2F2"   # light grey

STATUS_COLORS = {
    "Completed":          "D6EAD6",
    "Partially completed":"FFF3CD",
    "Not completed":      "F8D7DA",
    "Correctly excluded": "F2F2F2",
}

MOSCOW_COLORS = {
    "Must":   "1F3964",
    "Should": "2E6B9E",
    "Could":  "4A9CC7",
    "Won't":  "7F7F7F",
}

# ---------------------------------------------------------------------------
# XML helpers (shared with other build scripts)
# ---------------------------------------------------------------------------

def _shd(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    pPr.append(shd)


def _cell_shd(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _left_bar(paragraph, fill_hex, bar_color="1F3964"):
    _shd(paragraph, fill_hex)
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), bar_color)
    pBdr.append(left)
    pPr.append(pBdr)


def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = COL_HEADING
    return p


def add_body(doc, text):
    return doc.add_paragraph(text)


def add_note(doc, text, kind="note"):
    colours = {
        "note":    ("FFF3CD", "856404", "7B6000"),
        "info":    ("D1ECF1", "0C5460", "0C5460"),
        "success": ("D6EAD6", "1A5C2A", "1A5C2A"),
    }
    bg, left, txt_hex = colours.get(kind, colours["note"])
    prefix = {"note": "NOTE  ", "info": "INFO  ", "success": "SCOPE  "}.get(kind, "NOTE  ")
    p  = doc.add_paragraph()
    rb = p.add_run(prefix)
    rb.bold = True
    rb.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    rt = p.add_run(text)
    rt.font.color.rgb = RGBColor(*bytes.fromhex(txt_hex))
    _left_bar(p, bg, left)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


# ---------------------------------------------------------------------------
# Specialised requirement table
# ---------------------------------------------------------------------------

def _req_table(doc, headers, rows, col_widths, status_col=3, moscow_col=2):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"

    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _cell_shd(c, "1F3964")
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        # Determine row background from status column
        status_val = row_data[status_col] if status_col < len(row_data) else ""
        bg = "FFFFFF"
        for key, color in STATUS_COLORS.items():
            if status_val.startswith(key):
                bg = color
                break

        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            c.text = ""
            p = c.paragraphs[0]

            if ci == moscow_col:
                # MoSCoW badge: coloured bold text
                moscow = str(val)
                col_hex = MOSCOW_COLORS.get(moscow, "555555")
                r = p.add_run(moscow)
                r.bold = True
                r.font.color.rgb = RGBColor(*bytes.fromhex(col_hex))
                r.font.size = Pt(9.5)
            elif ci == status_col:
                # Status badge
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)
                if str(val).startswith("Completed") and not str(val).startswith("Partially"):
                    r.font.color.rgb = RGBColor(0x1A, 0x5C, 0x2A)
                    r.bold = True
                elif str(val).startswith("Partially"):
                    r.font.color.rgb = RGBColor(0x85, 0x64, 0x04)
                elif str(val).startswith("Not completed"):
                    r.font.color.rgb = RGBColor(0x84, 0x20, 0x29)
                else:
                    r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
            else:
                r = p.add_run(str(val))
                r.font.size = Pt(9.5)

            _cell_shd(c, bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)

    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build
# ---------------------------------------------------------------------------

FUNCTIONAL_REQS = [
    # (id, requirement, moscow, status, description)
    ("FR-01", "Finger landmark tracking using vision models",
     "Must", "Completed",
     "Real-time detection of all 21 hand landmarks for both hands using at least one open-source model."),
    ("FR-02", "MIDI ground truth capture (Casio CTK-2400)",
     "Must", "Completed",
     "Capture MIDI note_on events via USB with sub-millisecond timestamp precision."),
    ("FR-03", "Temporal synchronisation between MIDI and video",
     "Must", "Completed",
     "Align MIDI events and video frames to a shared perf_counter origin set at recording start."),
    ("FR-04", "Support multiple open-source models",
     "Must", "Partially completed",
     "MediaPipe Hands and OpenPose Hand integrated.  MoveNet excluded due to incompatible output format."),
    ("FR-05", "Camera-to-keyboard spatial calibration",
     "Must", "Completed",
     "Interactive homography-based polygon mask that warps key boundaries to match the camera's oblique view."),
    ("FR-06", "Data logging system",
     "Must", "Completed",
     "Per-session output: frames.csv (timestamps), MIDI JSONL, standard MIDI file, and session metadata JSON."),
    ("FR-07", "Keypress event analysis",
     "Must", "Completed",
     "Polygon containment test determines which key polygon a fingertip falls inside at each note_on event."),
    ("FR-08", "Quantitative performance metrics",
     "Must", "Completed",
     "MPJPE (horizontal-only), detection rate, and accuracy percentage computed per session and in aggregate."),
    ("FR-09", "Controlled variation of experimental conditions",
     "Must", "Completed",
     "Fitzpatrick skin type (II-V), lux category (Dim/Indoor), and hand size (7.5-12.5 cm) systematically logged."),
    ("FR-10", "Statistical comparison of models",
     "Must", "Completed",
     "Descriptive statistics (mean, SD) with grouped bar charts comparing MediaPipe and OpenPose outcomes."),
    ("FR-11", "Model selection justification",
     "Must", "Completed",
     "Literature review and methodology chapter justify model choices with reference to prior work."),
    ("FR-12", "Reproducible experimental protocol",
     "Should", "Completed",
     "Session protocol, hardware configuration, and analysis scripts documented and released."),
    ("FR-13", "Evaluation under additional models",
     "Could", "Not completed",
     "MoveNet and other models considered but excluded due to incompatible landmark output format."),
    ("FR-14", "Real-time visualisation overlay",
     "Should", "Completed",
     "live_preview.py shows hand skeleton, key polygons, MPJPE, and detection stats in real time."),
    ("FR-15", "Automatic occlusion event annotation",
     "Could", "Not completed",
     "Distinguishing true occlusion from model failure was descoped; counted as detection_fail."),
    ("FR-16", "3D hand pose estimation",
     "Won't", "Correctly excluded",
     "Single overhead camera cannot reconstruct 3D depth reliably; 2D horizontal MPJPE is sufficient."),
    ("FR-17", "Mobile / embedded deployment",
     "Won't", "Correctly excluded",
     "Research prototype; real-time laptop deployment meets all performance requirements."),
    ("FR-18", "Multi-camera / multi-view tracking",
     "Won't", "Correctly excluded",
     "Single fixed-overhead camera is the defined hardware configuration for this study."),
    ("FR-19", "Real-time audio feedback (gamification)",
     "Won't", "Correctly excluded",
     "Out of research scope; project focuses on measurement, not interactive music education."),
]

NON_FUNCTIONAL_REQS = [
    ("NFR-01", "Research-focused functional prototype",
     "Must", "Completed",
     "System functions as a data collection and analysis pipeline for academic research, not a consumer product."),
    ("NFR-02", "Fair and controlled comparison conditions",
     "Must", "Completed",
     "Identical hardware, camera position, and session protocol applied across all participants and both models."),
    ("NFR-03", "Python + open-source libraries only",
     "Must", "Completed",
     "Python 3.10, OpenCV, MediaPipe, Mido, NumPy, Pandas, Matplotlib.  No proprietary SDKs."),
    ("NFR-04", "Ethical compliance and participant protection",
     "Must", "Completed",
     "University ethics approval (Stages 1 and 2), signed informed consent forms, anonymised participant IDs."),
    ("NFR-05", "Real-time performance ≥ 15–20 FPS",
     "Should", "Completed",
     "Sustained 30 FPS on the host machine (MSI laptop, Intel i7-14th gen) during live_preview.py recording."),
    ("NFR-06", "Code quality, documentation, reproducibility",
     "Should", "Completed",
     "Organised script package, inline docstrings, INSTALLATION.md, User Guidelines, and this Requirement Catalogue."),
    ("NFR-07", "Basic robustness and error handling",
     "Could", "Completed",
     "Daemon-thread MIDI capture continues across hardware hiccups; graceful degradation when no MIDI port found."),
]


def build():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    # ── Title ────────────────────────────────────────────────────────────────
    t = doc.add_heading("Requirement Catalogue", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        r.font.color.rgb = COL_HEADING

    sub = doc.add_paragraph(
        "Comparative Evaluation of Vision-Based Finger Tracking for Piano Interaction"
    )
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.font.size   = Pt(12)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── 1. Purpose and scope ─────────────────────────────────────────────────
    add_heading(doc, "1.  Purpose and Scope", 1)
    add_body(doc,
        "This document catalogues all functional and non-functional requirements "
        "defined for the Piano Finger Tracking data collection and analysis system.  "
        "Requirements were established in the Initial Project Plan and are reproduced "
        "here in full with their MoSCoW priority and final completion status.")
    doc.add_paragraph()
    add_body(doc,
        "The system is a research prototype, not a consumer product.  Its purpose is "
        "to benchmark MediaPipe Hands and OpenPose Hand models on their ability to "
        "identify which finger strikes which piano key, under controlled variations "
        "of skin tone, lighting, and hand size.")
    doc.add_paragraph()
    add_note(doc,
        "MoSCoW key:  Must = critical to project success.  "
        "Should = high value but not blocking.  "
        "Could = desirable if time permits.  "
        "Won't = explicitly out of scope for this iteration.",
        kind="info")
    doc.add_paragraph()

    # ── 2. MoSCoW summary ────────────────────────────────────────────────────
    add_heading(doc, "2.  MoSCoW Priority Summary", 1)

    all_reqs = FUNCTIONAL_REQS + NON_FUNCTIONAL_REQS
    must_done    = sum(1 for r in all_reqs if r[2]=="Must"   and r[3].startswith("Completed"))
    must_total   = sum(1 for r in all_reqs if r[2]=="Must")
    should_done  = sum(1 for r in all_reqs if r[2]=="Should" and r[3].startswith("Completed"))
    should_total = sum(1 for r in all_reqs if r[2]=="Should")
    could_done   = sum(1 for r in all_reqs if r[2]=="Could"  and r[3].startswith("Completed"))
    could_total  = sum(1 for r in all_reqs if r[2]=="Could")
    wont_total   = sum(1 for r in all_reqs if r[2]=="Won't")

    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = "Table Grid"
    headers = ["Priority", "Total", "Completed", "Outcome"]
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _cell_shd(c, "1F3964")

    summary_rows = [
        ("Must",   str(must_total),   str(must_done),
         f"All {must_total} Must requirements completed" if must_done == must_total
         else f"{must_done}/{must_total} completed"),
        ("Should", str(should_total), str(should_done),
         f"All {should_total} Should requirements completed"),
        ("Could",  str(could_total),  str(could_done),
         f"{could_done}/{could_total} completed  ({could_total - could_done} descoped)"),
        ("Won't",  str(wont_total),   "N/A",
         f"All {wont_total} correctly excluded from scope"),
    ]
    moscow_bgs = {"Must": "D6EAD6", "Should": "E8F4FD", "Could": "FFF9E6", "Won't": "F2F2F2"}
    for ri, (priority, total, done, outcome) in enumerate(summary_rows):
        row = tbl.rows[ri + 1]
        bg  = moscow_bgs[priority]
        col = MOSCOW_COLORS[priority]
        vals = [priority, total, done, outcome]
        for ci, val in enumerate(vals):
            c = row.cells[ci]
            c.text = ""
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
            if ci == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(*bytes.fromhex(col))
            _cell_shd(c, bg)

    for row in tbl.rows:
        row.cells[0].width = Inches(0.9)
        row.cells[1].width = Inches(0.7)
        row.cells[2].width = Inches(1.0)
        row.cells[3].width = Inches(4.7)

    doc.add_paragraph()

    # ── 3. Functional requirements ───────────────────────────────────────────
    add_heading(doc, "3.  Functional Requirements", 1)
    add_body(doc,
        "Functional requirements define what the system must do.  "
        "Nineteen functional requirements were identified, covering data capture, "
        "model integration, calibration, analysis, and reporting.")
    doc.add_paragraph()

    _req_table(doc,
        ["ID", "Requirement", "MoSCoW", "Status", "Notes"],
        [(r[0], r[1], r[2], r[3], r[4]) for r in FUNCTIONAL_REQS],
        col_widths=[0.7, 2.3, 0.75, 1.2, 2.35],
        status_col=3, moscow_col=2,
    )

    # ── 3.1 Must requirements detail ─────────────────────────────────────────
    add_heading(doc, "3.1  Must Requirements  —  Detail", 2)
    must_reqs = [r for r in FUNCTIONAL_REQS if r[2] == "Must"]
    for req in must_reqs:
        p = doc.add_paragraph()
        p.add_run(f"{req[0]}  ").bold = True
        p.runs[-1].font.color.rgb = COL_HEADING
        p.add_run(req[1]).bold = True
        detail = doc.add_paragraph(req[4])
        detail.paragraph_format.left_indent = Cm(0.8)
        detail.paragraph_format.space_after = Pt(4)
        _shd(detail, "F5F7FA")

    doc.add_paragraph()

    # ── 3.2 Should/Could/Won't ───────────────────────────────────────────────
    add_heading(doc, "3.2  Should / Could / Won't Requirements  —  Detail", 2)
    other_reqs = [r for r in FUNCTIONAL_REQS if r[2] != "Must"]
    for req in other_reqs:
        p = doc.add_paragraph()
        p.add_run(f"{req[0]}  ").bold = True
        p.runs[-1].font.color.rgb = COL_HEADING
        p.add_run(req[1]).bold = True
        detail = doc.add_paragraph(req[4])
        detail.paragraph_format.left_indent = Cm(0.8)
        detail.paragraph_format.space_after = Pt(4)
        bg = "F5F7FA" if req[2] != "Won't" else "F2F2F2"
        _shd(detail, bg)

    doc.add_paragraph()

    # ── 4. Non-functional requirements ──────────────────────────────────────
    add_heading(doc, "4.  Non-Functional Requirements", 1)
    add_body(doc,
        "Non-functional requirements define how the system should perform — "
        "covering quality attributes such as performance, ethics, reproducibility, "
        "and technology constraints.")
    doc.add_paragraph()

    _req_table(doc,
        ["ID", "Requirement", "MoSCoW", "Status", "Notes"],
        [(r[0], r[1], r[2], r[3], r[4]) for r in NON_FUNCTIONAL_REQS],
        col_widths=[0.7, 2.3, 0.75, 1.2, 2.35],
        status_col=3, moscow_col=2,
    )

    add_heading(doc, "4.1  Non-Functional Requirements  —  Detail", 2)
    for req in NON_FUNCTIONAL_REQS:
        p = doc.add_paragraph()
        p.add_run(f"{req[0]}  ").bold = True
        p.runs[-1].font.color.rgb = COL_HEADING
        p.add_run(req[1]).bold = True
        detail = doc.add_paragraph(req[4])
        detail.paragraph_format.left_indent = Cm(0.8)
        detail.paragraph_format.space_after = Pt(4)
        _shd(detail, "F5F7FA")

    doc.add_paragraph()

    # ── 5. Requirements traceability ─────────────────────────────────────────
    add_heading(doc, "5.  Requirements Traceability", 1)
    add_body(doc,
        "The table below maps each Must and Should requirement to the script or "
        "system component that implements it.")
    doc.add_paragraph()

    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = "Table Grid"
    for i, h in enumerate(["Requirement ID", "Implemented in", "Verified by"]):
        c = tbl2.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.size = Pt(10)
        _cell_shd(c, "1F3964")

    traceability = [
        ("FR-01", "model_manager.py  /  live_preview.py",        "Live overlay + analyse.py output"),
        ("FR-02", "midi_recorder.MidiRecorder",                   "p###_midi.jsonl timestamps"),
        ("FR-03", "midi_recorder.py  (shared t0)",               "Frame/MIDI delta < 1 frame verified"),
        ("FR-04", "model_manager.MediaPipeHandModel + OpenPoseHandModel", "Both models run on all sessions"),
        ("FR-05", "key_calibration.KeyMask  (homography)",       "key_centers.json polygon vertices"),
        ("FR-06", "record.py  /  midi_recorder.py",              "data/raw/p###/ file listing"),
        ("FR-07", "analyse.py  (pointPolygonTest)",              "matched / detection_fail counts"),
        ("FR-08", "analyse.py  +  plot_results.py",              "data/processed/ JSON + charts"),
        ("FR-09", "fitzpatrick_detector.py  /  lux_calculator.py","p###_session.json fields"),
        ("FR-10", "plot_results.py  (bar charts, group means)",  "data/plots/ PNG files"),
        ("FR-11", "Report sections 3, 6, 8",                    "Documented in main report"),
        ("FR-12", "INSTALLATION.md  /  User Guidelines.docx",   "Released in repository"),
        ("FR-14", "live_preview.py",                             "Visual overlay + results panel"),
        ("NFR-01", "All scripts  (no UI beyond OpenCV windows)", "Prototype scope maintained"),
        ("NFR-02", "record.py setup screen  (identical protocol)", "Session metadata consistency"),
        ("NFR-03", "requirements.txt  +  piano_env/",           "All imports from open-source packages"),
        ("NFR-04", "Ethics docs  /  consent forms",             "Appendices C and D of main report"),
        ("NFR-05", "stats_collector.py  (FPS overlay)",         "30 FPS sustained in live_preview.py"),
        ("NFR-06", "scripts/ package  +  documentation files",  "Code review + docs suite"),
        ("NFR-07", "midi_recorder.py  (daemon thread)",         "MIDI capture continuity across errors"),
    ]
    for ri, (req_id, impl, verified) in enumerate(traceability):
        row = tbl2.add_row()
        bg  = "FFFFFF" if ri % 2 == 0 else "F5F7FA"
        vals = [req_id, impl, verified]
        for ci, val in enumerate(vals):
            c = row.cells[ci]
            c.text = ""
            r = c.paragraphs[0].add_run(val)
            r.font.size = Pt(9.5)
            if ci > 0:
                r.font.name = "Courier New"
            _cell_shd(c, bg)

    for row in tbl2.rows:
        row.cells[0].width = Inches(1.0)
        row.cells[1].width = Inches(3.2)
        row.cells[2].width = Inches(3.1)

    doc.add_paragraph()

    # ── 6. Out-of-scope items ────────────────────────────────────────────────
    add_heading(doc, "6.  Explicitly Out-of-Scope Items", 1)
    add_body(doc,
        "The following items were considered during project planning and deliberately "
        "excluded.  They are recorded here to prevent scope creep in future iterations.")
    doc.add_paragraph()

    wont_reqs = [r for r in FUNCTIONAL_REQS if r[2] == "Won't"]
    for req in wont_reqs:
        p = doc.add_paragraph()
        p.add_run(f"{req[0]}  ").bold = True
        p.runs[-1].font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)
        p.add_run(req[1]).bold = True
        detail = doc.add_paragraph(req[4])
        detail.paragraph_format.left_indent = Cm(0.8)
        detail.paragraph_format.space_after = Pt(4)
        _shd(detail, "F2F2F2")

    doc.add_paragraph()
    add_note(doc,
        "These Won't items represent deliberate design constraints, not failures.  "
        "They may be revisited in future work if hardware or scope changes.",
        kind="note")

    # ── Save ─────────────────────────────────────────────────────────────────
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print(f"Saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
